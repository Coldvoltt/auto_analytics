# tools.py

from crewai.tools import BaseTool
from langchain_experimental.utilities import PythonREPL
from pydantic import PrivateAttr
import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown2
from datetime import datetime


class PythonREPLCrewTool(BaseTool):
    name: str = "python_repl"
    description: str = (
        "A Python REPL tool. Use this to execute Python code. "
        "Make sure to use print(...) to produce output."
    )

    _repl: PythonREPL = PrivateAttr()

    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        object.__setattr__(self, "_repl", PythonREPL())

    def _run(self, code: str) -> str:
        try:
            result = self._repl.run(code)
            if result is None:
                return ""
            return str(result)
        except Exception as e:
            return f"Python REPL error: {e}"

    async def _arun(self, code: str) -> str:
        return self._run(code)


class MarkdownToDocxTool(BaseTool):
    name: str = "markdown_to_docx"
    description: str = (
        "Converts markdown content to a DOCX file. "
        "Provide the markdown content and output file path. "
        "Automatically handles images and formatting."
    )

    def _run(self, markdown_content: str, output_path: str = None) -> str:
        try:
            if output_path is None:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                output_path = f"outputs/reports/data_analysis_report_{timestamp}.docx"

            doc = Document()

            # Add document title
            doc.add_heading('Data Analysis Report', level=0)

            # Convert markdown to HTML for better parsing
            html = markdown2.markdown(markdown_content, extras=['tables', 'fenced-code-blocks'])

            # Process the markdown content line by line
            lines = markdown_content.split('\n')
            code_block = False
            code_content = []

            for line in lines:
                # Handle code blocks
                if line.startswith('```'):
                    if not code_block:
                        code_block = True
                        code_content = []
                    else:
                        code_block = False
                        # Add the code block
                        if code_content:
                            p = doc.add_paragraph()
                            p.style = 'No Spacing'
                            run = p.add_run('\n'.join(code_content))
                            run.font.name = 'Courier New'
                            run.font.size = Pt(9)
                    continue

                if code_block:
                    code_content.append(line)
                    continue

                # Handle headers
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('#### '):
                    doc.add_heading(line[5:], level=4)

                # Handle bullet points
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:], style='List Bullet')

                # Handle numbered lists
                elif re.match(r'^\d+\.\s', line):
                    text = re.sub(r'^\d+\.\s', '', line)
                    doc.add_paragraph(text, style='List Number')

                # Handle images
                elif re.match(r'^!\[.*\]\(.*\)$', line):
                    match = re.match(r'^!\[(.*)]\((.*)\)$', line)
                    if match:
                        alt_text = match.group(1)
                        img_path = match.group(2)

                        # Handle relative paths
                        if img_path.startswith('../charts/'):
                            img_path = img_path.replace('../charts/', 'outputs/charts/')

                        # Add image to document if it exists
                        if os.path.exists(img_path):
                            try:
                                p = doc.add_paragraph()
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run = p.add_run()
                                run.add_picture(img_path, width=Inches(6))

                                # Add caption
                                if alt_text:
                                    caption = doc.add_paragraph(alt_text)
                                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    caption.runs[0].font.size = Pt(10)
                                    caption.runs[0].italic = True
                            except Exception as img_error:
                                doc.add_paragraph(f"[Image: {alt_text} - Could not load: {img_error}]")
                        else:
                            doc.add_paragraph(f"[Image: {alt_text} - File not found: {img_path}]")

                # Handle tables (simplified)
                elif '|' in line and line.strip().startswith('|'):
                    # This is a simple table handler
                    cells = [cell.strip() for cell in line.split('|')[1:-1]]
                    if all(c == '-' or c == '---' for c in ''.join(cells).replace('-', '')):
                        continue  # Skip separator lines
                    # For now, just add as formatted text
                    doc.add_paragraph(line)

                # Handle empty lines
                elif line.strip() == '':
                    doc.add_paragraph()

                # Handle regular paragraphs with formatting
                else:
                    # Create paragraph
                    p = doc.add_paragraph()

                    # Handle bold and italic formatting
                    text = line

                    # Replace bold text
                    bold_pattern = r'\*\*(.*?)\*\*'
                    italic_pattern = r'\*(.*?)\*'
                    code_inline_pattern = r'`(.*?)`'

                    # Process inline formatting
                    current_pos = 0

                    # Find all formatting matches
                    all_matches = []

                    for match in re.finditer(bold_pattern, text):
                        all_matches.append(('bold', match))
                    for match in re.finditer(italic_pattern, text):
                        # Check it's not part of bold
                        if not any(m[1].start() <= match.start() < m[1].end() for m in all_matches if m[0] == 'bold'):
                            all_matches.append(('italic', match))
                    for match in re.finditer(code_inline_pattern, text):
                        all_matches.append(('code', match))

                    # Sort matches by position
                    all_matches.sort(key=lambda x: x[1].start())

                    # Process text with formatting
                    if all_matches:
                        for format_type, match in all_matches:
                            # Add text before the match
                            if match.start() > current_pos:
                                p.add_run(text[current_pos:match.start()])

                            # Add formatted text
                            if format_type == 'bold':
                                p.add_run(match.group(1)).bold = True
                            elif format_type == 'italic':
                                p.add_run(match.group(1)).italic = True
                            elif format_type == 'code':
                                run = p.add_run(match.group(1))
                                run.font.name = 'Courier New'

                            current_pos = match.end()

                        # Add remaining text
                        if current_pos < len(text):
                            p.add_run(text[current_pos:])
                    else:
                        p.add_run(text)

            # Ensure output directory exists
            os.makedirs(os.path.dirname(output_path), exist_ok=True)

            # Save the document
            doc.save(output_path)

            return f"Successfully converted markdown to DOCX: {output_path}"

        except Exception as e:
            return f"Error converting markdown to DOCX: {e}"

    async def _arun(self, markdown_content: str, output_path: str = None) -> str:
        return self._run(markdown_content, output_path)